#!/usr/bin/env python3
"""Черновики благодарственных писем (DOCX) на рабочий стол пользователя."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

DESKTOP = Path.home() / "Desktop"
OUT_DIR = DESKTOP / "APV_blagodarnosti_logistika"

RECIPIENT_LINES = (
    "Индивидуальному предпринимателю\n"
    "Махмадову Шарифхону Тагаймуродовичу\n"
    "140125, Московская обл., г. Люберцы, мкр. Новые Островцы (д. Островцы),\n"
    "ул. Лётчика Волчкова, д. 2, кв. 155"
)

LETTERS: list[dict[str, str]] = [
    {
        "fname": "Blagodarnost_sklad_Sofino.docx",
        "place": (
            "складского объекта в районе населённого пункта Софино "
            "(Раменский городской округ, Московская область)"
        ),
        "focus": (
            "операций приёмки и отгрузки, работ у ворот и на линии комплектации / погрузочно-разгрузочных работ"
        ),
        "sign": "И.П. Махмадов",
    },
    {
        "fname": "Blagodarnost_sklad_Britovo.docx",
        "place": "складского объекта в посёлке Бритово (Московская область)",
        "focus": "стабильного закрытия смен и дисциплины персонала на площадке в установленные окна",
        "sign": "И.П. Махмадов-Шариф",
    },
    {
        "fname": "Blagodarnost_sklad_Zhukovskiy.docx",
        "place": "складского объекта в городе Жуковский (Московская область)",
        "focus": (
            "согласования смен с транспортным и производственным графиком, оперативных замен в пиковые смены"
        ),
        "sign": "И.П. Махмадов",
    },
]


def _set_normal_font(doc: Document) -> None:
    sty = doc.styles["Normal"]
    sty.font.name = "Times New Roman"
    sty.font.size = Pt(12)


def _p(doc: Document, text: str, *, indent_mm: float = 0, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if indent_mm:
        p.paragraph_format.first_line_indent = Mm(indent_mm)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def build_letter(meta: dict[str, str]) -> Document:
    doc = Document()
    _set_normal_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("БЛАГОДАРНОСТЬ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    ad = doc.add_paragraph(RECIPIENT_LINES)
    ad.paragraph_format.space_after = Pt(12)
    for run in ad.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    _p(doc, "Уважаемый Шарифхон Тагаймуродович!", indent_mm=12.5)

    blocks = [
        (
            "Настоящим выражаем искреннюю признательность за результативное сотрудничество по закрытию "
            "складских смен и предоставлению персонала для нужд логистической компании на объекте "
            f"{meta['place']}."
        ),
        (
            "В ходе совместной работы специалисты вашей организации продемонстрировали понимание операционной "
            "специфики склада: соблюдение регламентов охраны труда и требований площадки, предсказуемую явку, "
            "прозрачный порядок замен и выстроенную коммуникацию с линейным руководством в периоды пиковой "
            "загрузки и повышенной сменности."
        ),
        (
            f"Отдельно отмечаем качество организации работ по {meta['focus']}, что позволило поддерживать "
            "приемлемый темп обработки грузопотоков и снизить операционные риски, связанные с неявкой и "
            "хаотичной подменой персонала."
        ),
        (
            "Рассчитываем на продолжение партнёрства и готовы рекомендовать вашу компанию как надёжного "
            "подрядчика по аутсорсингу складского персонала."
        ),
        "С уважением,",
    ]
    for b in blocks:
        _p(doc, b, indent_mm=12.5)

    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.paragraph_format.left_indent = Mm(70)
    sig.paragraph_format.space_before = Pt(8)
    r1 = sig.add_run("Руководитель группы складов логистики\n")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = sig.add_run("(зона ответственности: склады Софино, Бритово, Жуковский)\n\n")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r3 = sig.add_run(meta["sign"])
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(18)
    fr = foot.add_run(
        "Исх. № ______________   от «____» ____________________ 2026 г.\n\n"
        "Наименование логистической организации (отправителя): ___________________________________________\n\n"
        "М.П."
    )
    fr.font.size = Pt(10)
    fr.font.name = "Times New Roman"

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for meta in LETTERS:
        path = OUT_DIR / meta["fname"]
        build_letter(meta).save(path)
        print("Saved:", path)

    readme = OUT_DIR / "README.txt"
    readme.write_text(
        "Благодарственные письма — черновики для правки и подстановки реквизитов.\n\n"
        "Файлы .docx открываются в Microsoft Word, Pages или Google Docs.\n"
        "PDF: в Word — «Файл» → «Сохранить как» → PDF.\n\n"
        "Что подставить вручную: исходящий номер, дата, полное наименование и реквизиты логистической "
        "организации-отправителя, при необходимости — скан печати (М.П.).\n\n"
        "Подписи в черновиках:\n"
        "  • Софино — И.П. Махмадов\n"
        "  • Бритово — И.П. Махмадов-Шариф\n"
        "  • Жуковский — И.П. Махмадов\n\n"
        "Если нужна третья строка с расшифровкой инициалов — добавьте под подписью от руки или второй строкой.\n",
        encoding="utf-8",
    )
    print("README:", readme)
    print("Папка:", OUT_DIR)


if __name__ == "__main__":
    main()
